'''
    Forward Kinematics Table Generator:

    This code is designed to generate a Forward Kinematics Table based on the provided Denavit-Hartenberg (DH) parameters.
    
    Authors: Passias Konstantinos, Michael Chatzimichael

    Instructions:
    1. Ensure that DH parameters are correctly defined for your robot.
    2. Run the code to generate the Forward Kinematics Table.

---------------------------------------------------------------------------------------

    Inverse Kinematics Validation:

    After obtaining the Forward Kinematics Table, use the following steps to validate Inverse Kinematics:
    
    1. Calculate the Inverse Kinematics equations for your robot.
    2. Substitute theta values with 0 to validate the results.
    
    Authors: Passias Konstantinos, Michael Chatzimichael

    Note: Make sure to double-check DH parameters and equations for accurate results.

---------------------------------------------------------------------------------------

    Feel free to modify the code and comments based on your specific robot configuration.

---------------------------------------------------------------------------------------
    
    Probs. It prints twice the Trasnformation Matrices because of Tlink function

'''


from sympy import *
from sympy.matrices import Matrix
import docx
from docx.shared import Pt
import math

doc = docx.Document()
# Settings
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(8)

scale =1
transformation_matrices = []


def ForwardKinematics():
    d1, theta1, theta2, theta3, L2, L3 = symbols('d1, theta1, theta2, theta3, L2, L3')
    theta4, theta5 = symbols('theta4, theta5')
    a1, a2, a3, d3 = symbols('a1, a2, a3, d3')
    

    
    #Validation for Inverse (comment if not needed)
    # theta1=math.radians(122)
    # theta2=math.radians(-71)
    # theta3=math.radians(0)
    # scale = 1000
    # a1 = 215.6 / scale
    # d1 = 298.7 / scale
    # a2 = 274.51 / scale
    # a3 = 71.3 / scale
    # #end validation for Inverse

    # a | α | d | θ 
    DH4 = Matrix([
        [0,pi/2,d1,theta1],
        [a1,-pi/2,0,theta2],
        [a2,0,0.071,theta3]
    ])

    DH = DH4

    TG0 = TLink(DH.row(0))
    T01 = TLink(DH.row(1))
    T12 = TLink(DH.row(2))

    TGe = TG0 * T01 * T12 

    TGe = simplify(TGe)
    

    for i in range(3):  # Adjust the range based on the number of matrices
        T_result = TLink(DH.row(i))
        transformation_matrices.append(T_result)
    
    transformation_matrices.append(TGe)

    return TGe


def TLink(DH):
    # DH parameters
    a, alpha, d, theta = DH

    # Initialize a symbolic T matrix
    T = Matrix([
        [cos(theta), -sin(theta), 0, a],
        [sin(theta)*cos(alpha), cos(theta)*cos(alpha), -sin(alpha), -sin(alpha)*d],
        [sin(theta)*sin(alpha), cos(theta)*sin(alpha), cos(alpha), cos(alpha)*d],
        [0, 0, 0, 1]
    ])

    # print("Matrix T:")
    # pprint(T)

    return T


def create_table(doc, matrix):
    table = doc.add_table(rows=matrix.shape[0], cols=matrix.shape[1], style='Table Grid')

    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            cell = table.cell(i, j)
            cell.text = pretty(matrix[i, j], use_unicode=True)

def save_to_word(result_matrices):
    total_matrices = len(result_matrices)
    
    for i, matrix in enumerate(result_matrices):
        if i == total_matrices - 1:
            doc.add_heading(f'Final Transformation Matrix TGe', level=1)
        else:
            doc.add_heading(f'Transformation Matrix T{i}-{i+1}', level=1)
        create_table(doc, matrix)

    doc.save('forward_kinematics_result.docx')
    print("Word document saved with Forward Kinematics result.")

def ask_yes_no_question():
    while True:
        user_input = input("Do you want to save results? (yes/no): ").lower()
        
        if user_input == 'yes':
            print("You chose 'yes'. Proceeding...")
            break
        elif user_input == 'no':
            print("You chose 'no'. Exiting...")
            break
        else:
            print("Invalid input. Please enter 'yes' or 'no'.")
    return user_input

if __name__ == "__main__":
    TGe_result = ForwardKinematics()
    

    print("Result of Forward Kinematics:")
    pprint(TGe_result)


    user_input = ask_yes_no_question()
    if user_input == 'yes':
        save_to_word(transformation_matrices)
